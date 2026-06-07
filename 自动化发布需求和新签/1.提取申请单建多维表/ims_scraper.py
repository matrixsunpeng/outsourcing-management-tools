"""
IMS 自动化 — 导出Excel → 筛选新签 → 逐条查询弹窗 → 提取14个字段
"""
import os
import re
import json
import time
import tempfile
import requests
from pathlib import Path
from urllib.parse import urljoin

from playwright.sync_api import sync_playwright
from dotenv import load_dotenv

load_dotenv()

IMS_URL = "https://ims.asiainfo.com/AIOMS/Jsp/main.jsp"
IMS_USERNAME = os.getenv("IMS_USERNAME", "")
IMS_PASSWORD = os.getenv("IMS_PASSWORD", "")

SBU_VALUE = "185"          # 亚信科技CMB flexValue
COOP_TYPE_INDEX = 2        # 技术合作-||(人员类)
APP_STATE_ITEMCODE = "40"  # 审批流程结束 itemCode

DOWNLOAD_DIR = Path(__file__).parent / "downloads"
EXCEL_PATH = Path(__file__).parent / "exported_data.xlsx"


# 从弹窗提取的14个字段
POPUP_FIELDS = [
    "合作申请单编号", "事业部/SBU", "统计区域", "申请人",
    "签约性质", "申请日期", "资源池代码", "审批状态",
    "技术合作人员数量", "预计技术合作时成本",
    "技术合作服务周期开始日期", "技术合作服务周期结束日期",
    "技术合作需求明细", "备注",
]


class IMSScraper:
    def __init__(self, headless=False):
        self.headless = headless
        self.browser = None
        self.page = None
        self.playwright = None
        self._query_frame = None
        DOWNLOAD_DIR.mkdir(exist_ok=True)

    def start(self):
        self.playwright = sync_playwright().start()
        self.browser = self.playwright.chromium.launch(headless=self.headless)
        context = self.browser.new_context(
            viewport={"width": 1366, "height": 768},
            accept_downloads=True,
        )
        self.page = context.new_page()
        self.page.set_default_timeout(30000)

    # ==================== 登录 ====================
    def login(self):
        print("[IMS] 登录...")
        self.page.goto(IMS_URL, wait_until="networkidle", timeout=30000)
        if "sso.asiainfo.com" in self.page.url:
            self.page.fill("input[name='username']", IMS_USERNAME)
            self.page.fill("input[name='password']", IMS_PASSWORD)
            try:
                cb = self.page.locator("input#agreement")
                if cb.is_visible(timeout=1000) and not cb.is_checked():
                    cb.check()
            except Exception:
                pass
            self.page.click("input[name='submit']")
            self.page.wait_for_timeout(8000)
        if "login" in self.page.url.lower() or "sso" in self.page.url:
            raise RuntimeError("登录失败")
        print("[IMS] 登录完成")

    # ==================== 导航 ====================
    def navigate_to_query(self):
        print("[IMS] 导航到申请单...")
        self.page.wait_for_timeout(2000)
        for n in self.page.locator("span.mini-tree-nodetext").all():
            try:
                if n.inner_text().strip() == "外包数据查询":
                    n.locator("..").locator("..").dblclick()
                    self.page.wait_for_timeout(1500)
                    break
            except Exception:
                continue
        for n in self.page.locator("span.mini-tree-nodetext").all():
            try:
                if n.inner_text().strip() == "申请单":
                    n.click()
                    self.page.wait_for_timeout(3000)
                    break
            except Exception:
                continue
        for f in self.page.frames:
            if "queryManager" in f.url:
                self._query_frame = f
                break
        if not self._query_frame:
            raise RuntimeError("未找到查询 iframe")
        print("[IMS] 已进入申请单查询页面")

    # ==================== 导出 Excel ====================
    def export_excel(self, date_start, date_end):
        """设置查询条件 → 点查询 → 导出 Excel"""
        tf = self._query_frame
        print(f"[IMS] 设置查询条件: {date_start}~{date_end}")

        # 清理旧下载文件
        for f in DOWNLOAD_DIR.glob("*.xls*"):
            try:
                f.unlink()
            except Exception:
                pass

        # 设置查询条件
        tf.evaluate(f"""
            (function() {{
                mini.get("p_apply_begin_date").setValue("{date_start}");
                mini.get("p_apply_end_date").setValue("{date_end}");
                mini.get("p_coop_type").select({COOP_TYPE_INDEX});
                mini.get("p_app_state").select(4);
                mini.get("p_sbu_id").setValue("{SBU_VALUE}");
            }})()
        """)

        # 先点"查询"
        print("[IMS] 点击查询...")
        tf.evaluate('''
            (function() {
                var links = document.querySelectorAll("a.g_a");
                for (var i = 0; i < links.length; i++) {
                    if (links[i].innerText.trim() === "查询") {
                        links[i].click();
                        return;
                    }
                }
            })()
        ''')
        self.page.wait_for_timeout(3000)

        total = tf.evaluate('mini.get("datagrid1").getTotalCount()')
        print(f"[IMS] 查询结果: {total} 条")

        if total == 0:
            print("[IMS] 无结果，跳过导出")
            return 0

        # 使用 expect_download 处理下载
        with self.page.expect_download() as download_info:
            tf.evaluate('''
                (function() {
                    var links = document.querySelectorAll("a.g_a");
                    for (var i = 0; i < links.length; i++) {
                        if (links[i].innerText.indexOf("导出Excel") >= 0) {
                            links[i].click();
                            return;
                        }
                    }
                })()
            ''')

        download = download_info.value
        save_path = str(DOWNLOAD_DIR / download.suggested_filename)
        download.save_as(save_path)
        try:
            download.delete()
        except Exception:
            pass
        print(f"[IMS] 下载完成: {download.suggested_filename}")

        return total

    # ==================== 解析 Excel ====================
    def parse_and_filter_excel(self):
        """解析导出的 Excel（兼容 .xls / .xlsx），筛选 签约性质=='新签'"""
        # 查找下载的 Excel 文件
        excel_files = list(DOWNLOAD_DIR.glob("*.xls*"))
        if not excel_files:
            print("[IMS] 未找到导出的 Excel 文件")
            return []

        # 取最新的
        excel_path = max(excel_files, key=lambda p: p.stat().st_mtime)
        print(f"[IMS] 读取 Excel: {excel_path.name}")

        # 判断文件格式
        fname = excel_path.name.lower()
        if fname.endswith(".xls") or self._is_ole2(excel_path):
            records = self._parse_xls(excel_path)
        else:
            records = self._parse_xlsx(excel_path)

        return records

    def _is_ole2(self, path):
        """检查是否是 OLE2 格式（旧 .xls）"""
        with open(path, "rb") as f:
            return f.read(4) == b"\xd0\xcf\x11\xe0"

    def _parse_xls(self, path):
        """用 xlrd 解析 .xls"""
        import xlrd
        wb = xlrd.open_workbook(path)
        ws = wb.sheet_by_index(0)

        # 第2行(0-indexed row 1)是标题
        headers = [str(ws.cell_value(1, c)).strip() for c in range(ws.ncols)]

        sign_col = None
        for i, h in enumerate(headers):
            if "签约性质" in h:
                sign_col = i
                break
        if sign_col is None:
            print("[IMS] 未找到「签约性质」列")
            return []

        records = []
        for r in range(2, ws.nrows):
            sign = str(ws.cell_value(r, sign_col)).strip()
            if "新签" in sign:
                rec = {}
                for c, h in enumerate(headers):
                    val = ws.cell_value(r, c)
                    rec[h] = str(val).strip() if val else ""
                records.append(rec)

        print(f"[IMS] Excel 共 {ws.nrows - 2} 条, 筛选「新签」{len(records)} 条")
        return records

    def _parse_xlsx(self, path):
        """用 openpyxl 解析 .xlsx"""
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        ws = wb.active
        headers = []
        for row in ws.iter_rows(values_only=True, min_row=2, max_row=2):
            headers = [str(c).strip() if c else "" for c in row]

        sign_col = None
        for i, h in enumerate(headers):
            if "签约性质" in h:
                sign_col = i
                break
        if sign_col is None:
            return []

        records = []
        for row in ws.iter_rows(values_only=True, min_row=3):
            sign = str(row[sign_col]).strip() if sign_col < len(row) else ""
            if "新签" in sign:
                rec = {}
                for i, h in enumerate(headers):
                    val = row[i] if i < len(row) else ""
                    rec[h] = str(val).strip() if val else ""
                records.append(rec)

        print(f"[IMS] Excel 共 {ws.max_row - 2} 条, 筛选「新签」{len(records)} 条")
        wb.close()
        return records

    # ==================== 逐条查询弹窗详情 ====================
    def search_and_extract(self, application_code):
        """用合作申请单编号搜索，点击流水号打开详情，提取14个字段"""
        print(f"[IMS] 搜索: {application_code}")

        # 重新导航到查询页面
        self._reenter_query_page()

        tf = self._query_frame

        # 1) 设置宽日期范围并仅按合作申请单编号搜索
        tf.evaluate(f"""
            (function() {{
                mini.get("p_apply_begin_date").setValue("2026-01-01");
                mini.get("p_apply_end_date").setValue("2026-12-31");
                mini.get("p_coop_type").setValue("");
                mini.get("p_app_state").setValue("");
                mini.get("p_sbu_id").setValue("");
                mini.get("p_application_code").setValue("{application_code}");
            }})()
        """)

        # 2) 点击查询
        tf.evaluate('document.querySelector("a.g_a").click()')
        self.page.wait_for_timeout(3000)

        # 3) 检查结果
        total = tf.evaluate('mini.get("datagrid1").getTotalCount()')
        if total == 0:
            print(f"  [IMS] 未找到: {application_code}")
            return {}

        # 4) 点击流水号链接
        tf.evaluate("""
            (function() {
                var rows = document.querySelectorAll("tr.mini-grid-row");
                if (rows.length > 0) {
                    var a = rows[0].querySelector("a");
                    if (a) {
                        var href = a.getAttribute("href");
                        if (href) {
                            // 强制完整导航而非缓存切换
                            window.open(href, "_blank");
                            return "new_tab";
                        }
                    }
                }
                return "no_link";
            })()
        """)
        self.page.wait_for_timeout(5000)

        # 5a) 检查新标签页
        detail_text = ""
        attachments = []
        pages = self.page.context.pages
        if len(pages) > 1:
            detail_page = pages[-1]
            detail_page.bring_to_front()
            detail_page.wait_for_load_state("networkidle", timeout=15000)
            detail_text = detail_page.locator("body").inner_text()
            # 下载附件——通过 exportSingle 函数
            attachment_content = self._download_detail_attachments(detail_page)
            detail_page.close()
        else:
            # 5b) 尝试在 iframe 中查找
            attachment_content = ""
            for f in self.page.frames:
                if "applicationNoPro" in f.url:
                    try:
                        detail_text = f.evaluate("document.body.innerText")
                        attachment_content = self._download_detail_attachments(f)
                    except Exception:
                        pass
                    break

        if not detail_text:
            print("  [IMS] 未找到详情内容")
            return {}

        # 6) 解析详情字段
        detail = self._parse_detail_text(detail_text)

        # 7) 附件内容覆盖文本值
        if attachment_content:
            detail["技术合作需求明细"] = attachment_content
            print(f"  [IMS] 附件内容: {len(attachment_content)} 字符")

        print(f"  [IMS] 已提取: {list(detail.keys())}")
        return detail

    def _reenter_query_page(self):
        """重新进入查询页面——直接点击申请单"""
        # 直接点击"申请单"重新加载
        for n in self.page.locator("span.mini-tree-nodetext").all():
            try:
                if n.inner_text().strip() == "申请单":
                    n.click()
                    self.page.wait_for_timeout(4000)
                    break
            except Exception:
                continue

        # 重新获取 query frame
        for f in self.page.frames:
            if "queryManager" in f.url and "applicationNoPro" not in f.url:
                self._query_frame = f
                return

        raise RuntimeError("无法重新进入查询页面")

    def _parse_detail_text(self, text):
        """从详情文本解析14个字段"""
        result = {}

        # 字段解析规则: (页面标签, 输出字段名)
        rules = [
            ("合作申请单编号:", "合作申请单编号"),
            ("申请部门:", "事业部/SBU"),
            ("区域-大区:", "统计区域"),
            ("申请人:", "申请人"),
            ("签约性质:", "签约性质"),
            ("申请日期:", "申请日期"),
            ("资源池代码:", "资源池代码"),
            ("技术合作人员数量:", "技术合作人员数量"),
            ("预计技术合作时成本:", "预计技术合作时成本"),
            ("技术合作服务周期开始日期:", "技术合作服务周期开始日期"),
            ("技术合作服务周期结束日期:", "技术合作服务周期结束日期"),
            ("技术合作需求明细:", "技术合作需求明细"),
            ("备注:", "备注"),
        ]

        lines = text.split("\n")

        # 方法1: Tab分隔格式（紧凑模式）
        for line in lines:
            line = line.strip()
            if not line:
                continue
            parts = line.split("\t")
            for i, part in enumerate(parts):
                part = part.strip()
                if not part:
                    continue
                for label, field_name in rules:
                    if part == label or part.rstrip(":") == label.rstrip(":"):
                        if i + 1 < len(parts):
                            val = parts[i + 1].strip()
                            if val and len(val) < 200 and ":" not in val:
                                result[field_name] = val
                        break
                    if part.startswith(label) and len(part) > len(label):
                        val = part[len(label):].strip()
                        if val and len(val) < 200:
                            result[field_name] = val
                        break

        # 方法2: 分行格式——标签独占一行，值在后续行
        skip_prefixes = ("{", ".", "My JSP", "padding", "margin", "border",
                         "word-break", "background", "empty-cells", "font-",
                         "height", "width", "}", "tableMain")
        for i, line in enumerate(lines):
            line = line.strip()
            for label, field_name in rules:
                if field_name in result:
                    continue
                if line in (label, label.rstrip(":")):
                    for j in range(i + 1, min(i + 8, len(lines))):
                        nl = lines[j].strip()
                        if not nl or any(nl.startswith(p) for p in skip_prefixes):
                            continue
                        if any(nl in (l, l.rstrip(":")) for l, _ in rules):
                            break
                        if len(nl) < 200:
                            result[field_name] = nl
                            break
                    break

        if "审批状态" not in result:
            result["审批状态"] = "审批流程结束"

        return result

    def _find_attachments_in_page(self, page):
        """查找页面中的 Word/Excel 附件链接"""
        try:
            hrefs = page.evaluate("""
                (function() {
                    var links = document.querySelectorAll('a[href]');
                    var result = [];
                    for (var i = 0; i < links.length; i++) {
                        var h = links[i].getAttribute('href') || '';
                        if (h.match(/\\.(docx?|xlsx?)$/i)) result.push(h);
                    }
                    return JSON.stringify(result);
                })()
            """)
            return json.loads(hrefs)
        except Exception:
            return []

    def _download_attachments(self, urls):
        if not urls:
            return ""
        texts = []
        for url in urls:
            content = self._download_and_read(url)
            if content:
                texts.append(content)
        return "\n\n".join(texts)

    def _download_detail_attachments(self, frame_or_page):
        """在详情页中查找并下载 exportSingle 附件，返回文件内容"""
        try:
            # 查找所有带 exportSingle onclick 的链接
            file_ids = frame_or_page.evaluate("""
                (function() {
                    var links = document.querySelectorAll('a');
                    var result = [];
                    for (var i = 0; i < links.length; i++) {
                        var oc = links[i].getAttribute('onclick') || '';
                        var match = oc.match(/exportSingle\\('([^']+)'\\)/);
                        if (match) {
                            result.push({id: match[1], name: links[i].innerText.trim()});
                        }
                    }
                    return JSON.stringify(result);
                })()
            """)
            file_infos = json.loads(file_ids)
            if not file_infos:
                return ""

            all_content = []
            for fi in file_infos:
                print(f"  [附件] 下载: {fi['name']} (id={fi['id']})")
                # 用 requests 直接下载（exportSingle 是表单提交）
                content = self._download_via_http(fi['id'])
                if content:
                    all_content.append(content)

            return "\n\n---\n\n".join(all_content)

        except Exception as e:
            print(f"  [附件] 下载失败: {e}")
            return ""

    def _download_via_http(self, file_id):
        """通过 HTTP 直接下载附件并读取内容"""
        try:
            cookies = self.page.context.cookies()
            session = requests.Session()
            for c in cookies:
                session.cookies.set(c["name"], c["value"], domain=c.get("domain", ""))

            url = f"https://ims.asiainfo.com/AIOMS/uploadFile_downLoadFile.action?appAttachmentId={file_id}"
            resp = session.get(url, timeout=30)

            if resp.status_code != 200:
                print(f"  [附件] HTTP {resp.status_code}")
                return ""

            # 从 Content-Disposition 获取文件名
            cd = resp.headers.get("Content-Disposition", "")
            fname = ""
            if "filename=" in cd:
                fname = cd.split("filename=")[-1].strip('"')

            suffix = Path(fname).suffix.lower() if fname else ".tmp"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
                f.write(resp.content)
                tmp_path = f.name

            content = self._read_file_content(tmp_path)
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
            return content

        except Exception as e:
            print(f"  [附件] HTTP下载失败: {e}")
            return ""

    def _read_file_content(self, path):
        """读取文件全文内容"""
        fname = Path(path).name.lower()
        try:
            if fname.endswith('.xlsx'):
                import openpyxl
                wb = openpyxl.load_workbook(path, data_only=True)
                lines = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        row_text = " | ".join(str(c) for c in row if c is not None)
                        if row_text.strip():
                            lines.append(row_text)
                return "\n".join(lines)
            elif fname.endswith('.xls'):
                import xlrd
                wb = xlrd.open_workbook(path)
                ws = wb.sheet_by_index(0)
                lines = []
                for r in range(ws.nrows):
                    row_vals = [str(ws.cell_value(r, c)) for c in range(ws.ncols) if ws.cell_value(r, c) != '']
                    if row_vals:
                        lines.append(" | ".join(row_vals))
                return "\n".join(lines)
            elif fname.endswith('.docx'):
                import docx
                doc = docx.Document(path)
                parts = []
                # 读取段落
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text.strip())
                # 读取表格
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(
                            cell.text.strip() for cell in row.cells if cell.text.strip()
                        )
                        if row_text.strip():
                            parts.append(row_text)
                return "\n".join(parts)
            elif fname.endswith('.doc'):
                # Old .doc - try reading as text
                with open(path, 'r', encoding='utf-8', errors='replace') as f:
                    return f.read()[:10000]
        except Exception as e:
            print(f"  [附件] 读取失败: {e}")
        return ""

    def _download_and_read(self, href):
        try:
            cookies = self.page.context.cookies()
            session = requests.Session()
            for c in cookies:
                session.cookies.set(c["name"], c["value"], domain=c.get("domain", ""))
            url = href if href.startswith("http") else urljoin(self.page.url, href)
            resp = session.get(url, timeout=30)
            if resp.status_code != 200:
                return ""

            suffix = Path(href).suffix.lower()
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
                f.write(resp.content)
                tmp_path = f.name

            content = self._read_file_content(tmp_path)
            os.unlink(tmp_path)
            return content.strip() if content else ""
        except Exception as e:
            print(f"  [附件] 下载失败: {e}")
            return ""

    def close(self):
        if self.browser:
            self.browser.close()
        if self.playwright:
            self.playwright.stop()


def run_full_extraction(date_start, date_end):
    """完整提取流程"""
    scraper = IMSScraper(headless=False)
    try:
        scraper.start()
        scraper.login()
        scraper.navigate_to_query()

        # Step 1: 导出并筛选
        scraper.export_excel(date_start, date_end)
        records = scraper.parse_and_filter_excel()
        if not records:
            print("[IMS] 无新签记录")
            return []

        # Step 2: 逐条查询弹窗详情
        details = []
        for i, rec in enumerate(records):
            app_code = rec.get("合作申请单编号", "")
            if not app_code:
                continue

            print(f"\n[IMS] === {i+1}/{len(records)} ===")
            detail = scraper.search_and_extract(app_code)

            if detail:
                # 确保合作申请单编号正确
                if not detail.get("合作申请单编号"):
                    detail["合作申请单编号"] = app_code
                details.append(detail)
                print(f"  已提取字段: {[k for k in detail if not k.startswith('_')]}")

        return details
    finally:
        scraper.close()
